"""
GhostStrike Report Studio - Custom Company Templates
=========================================================
Item 29: `gs report template fill <company-template.docx>` -- a consultant
supplies their own .docx with placeholders like {{client_name}},
{{findings}}; GhostStrike fills them in with real engagement data.

Two placeholder kinds:
  - Scalar (client_name, operator, environment, scope, executive_summary,
    methodology, remediation_summary): the whole paragraph's text is
    replaced. Word routinely splits one visible sentence across multiple
    `<w:r>` runs (spell-check boundaries, mid-sentence formatting changes,
    autocorrect markers) even when the user never intended separate runs,
    so a naive per-run string search misses most real templates -- this
    reads each paragraph's *joined* text to find the placeholder, then
    rewrites the paragraph's runs as a single run with the substituted
    text. Formatting finer than "whatever the first run had" is not
    preserved for that paragraph -- an accepted simplification, the same
    one docxtpl-style tools make, rather than pulling in a heavier
    templating dependency for it.
  - Block (findings, critical_findings, attack_paths): these render a
    real table, not a text string, so the entire placeholder paragraph is
    replaced with a table inserted at that exact position in the
    document -- python-docx has no direct "insert element here" API for
    this, so it's done via the underlying paragraph XML element's
    addnext(), a documented python-docx pattern for this exact need.

© 2026 Fouad Ailabouni. All rights reserved.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List

from .data_sources import ReportData

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")

_BLOCK_PLACEHOLDERS = {"findings", "critical_findings", "attack_paths"}


def _scalar_values(data: ReportData) -> Dict[str, str]:
    eng = data.engagement or {}
    counts = data.severity_counts
    summary_bits = [f"{sev}: {counts[sev]}" for sev in ("CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO") if counts.get(sev)]

    modules_run = sorted({r.get("module", "") for r in data.runs if r.get("module")})
    methodology = ("Testing was conducted using the following GhostStrike modules: "
                   + ", ".join(modules_run) + ".") if modules_run else \
        "No module runs are recorded in the run ledger for this engagement."

    scope_text = ""
    scope_file = eng.get("scope_file", "")
    if scope_file and Path(scope_file).exists():
        scope_text = Path(scope_file).read_text(encoding="utf-8", errors="replace").strip()
    if not scope_text:
        scope_text = "Scope not recorded (no scope file set for this engagement)."

    if counts.get("CRITICAL"):
        exec_summary = (f"This assessment identified {counts['CRITICAL']} critical and {counts.get('HIGH', 0)} "
                         f"high-severity finding(s) requiring immediate remediation.")
    elif counts.get("HIGH"):
        exec_summary = f"This assessment identified {counts['HIGH']} high-severity finding(s) requiring remediation."
    elif data.findings:
        exec_summary = f"This assessment identified {len(data.findings)} finding(s), none critical or high severity."
    else:
        exec_summary = "No findings were recorded for this engagement."

    remediated = [f for f in data.findings if f.get("status") == "fixed"]
    still_open = [f for f in data.findings if f.get("status") not in ("fixed", "accepted_risk")]
    remediation_summary = (f"{len(remediated)} finding(s) have been remediated and verified. "
                            f"{len(still_open)} finding(s) remain open.")

    return {
        "client_name": eng.get("client", "") or "N/A",
        "operator": eng.get("operator", "") or "N/A",
        "environment": eng.get("environment", "") or "N/A",
        "engagement_id": data.engagement_id,
        "scope": scope_text,
        "executive_summary": exec_summary,
        "methodology": methodology,
        "remediation_summary": remediation_summary,
        "severity_summary": ", ".join(summary_bits) or "No findings recorded.",
    }


def _set_paragraph_text(paragraph, text: str) -> None:
    """Collapse a paragraph to a single run carrying `text`, preserving
    the first run's formatting (bold/italic/font) as the paragraph's own,
    since that's the only formatting choice a template author actually
    controls per-placeholder in practice."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _reaches_crown_jewel(finding: dict, ghost_scores: Dict[str, dict]) -> bool:
    gs = ghost_scores.get(finding.get("finding_id", finding.get("id", "")))
    return bool(gs and gs["factors"].get("crown_jewel_reachable"))


def _insert_table_at(paragraph, headers: List[str], rows: List[List[str]]):
    """Inserts a real table immediately after `paragraph` in document
    order, then removes the (now-empty) placeholder paragraph -- the
    python-docx pattern for positional insertion, since add_table() alone
    always appends at the document's end."""
    doc = paragraph.part.document
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_values in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row_values):
            cells[i].text = str(v)

    paragraph._p.addnext(table._tbl)
    paragraph._p.getparent().remove(paragraph._p)


def fill_template(template_path: Path, data: ReportData, out_path: Path,
                   crown_jewel_hosts: List[str] | None = None) -> Path:
    import docx
    import sys as _sys

    lib_dir = Path(__file__).resolve().parent.parent.parent / "bash_scripts_for_pentest" / "lib"
    if str(lib_dir) not in _sys.path:
        _sys.path.insert(0, str(lib_dir))
    import ghost_score as gscore

    doc = docx.Document(str(template_path))
    scalars = _scalar_values(data)

    ghost_scores: Dict[str, dict] = {}
    if crown_jewel_hosts:
        try:
            graph = gscore.agb.build_graph(data.engagement_id)
            for f in data.findings:
                scored = gscore.score_finding(f, graph, crown_jewel_hosts)
                ghost_scores[scored["finding_id"]] = scored
        except Exception:
            ghost_scores = {}

    # Two passes: block placeholders (structural, paragraph-replacing)
    # first, since they mutate the paragraph list; scalar text
    # substitution second, over whatever paragraphs remain.
    for paragraph in list(doc.paragraphs):
        text = paragraph.text
        match = _PLACEHOLDER_RE.fullmatch(text.strip())
        if not match or match.group(1) not in _BLOCK_PLACEHOLDERS:
            continue
        kind = match.group(1)

        if kind == "findings":
            rows = [[f.get("severity", ""), f.get("title", ""), f.get("target", {}).get("host", ""),
                     f.get("status", "open")] for f in data.findings_sorted()]
            _insert_table_at(paragraph, ["Severity", "Title", "Host", "Status"], rows)
        elif kind == "critical_findings":
            crit = [f for f in data.findings_sorted() if f.get("severity") in ("CRITICAL", "HIGH")]
            rows = [[f.get("severity", ""), f.get("title", ""), f.get("remediation", "")] for f in crit]
            _insert_table_at(paragraph, ["Severity", "Title", "Remediation"], rows)
        elif kind == "attack_paths":
            reaching = [f for f in data.findings_sorted() if _reaches_crown_jewel(f, ghost_scores)]
            rows = [[f.get("title", ""), f.get("target", {}).get("host", ""),
                     ghost_scores[f.get("finding_id", f.get("id", ""))]["ghost_score"]] for f in reaching]
            _insert_table_at(paragraph, ["Finding", "Host", "GhostScore"], rows)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "{{" not in text:
            continue
        new_text = _PLACEHOLDER_RE.sub(lambda m: scalars.get(m.group(1), m.group(0)), text)
        if new_text != text:
            _set_paragraph_text(paragraph, new_text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


_SCALAR_PLACEHOLDERS = {
    "client_name", "operator", "environment", "engagement_id", "scope",
    "executive_summary", "methodology", "remediation_summary", "severity_summary",
}


def known_placeholders() -> List[str]:
    return sorted(_SCALAR_PLACEHOLDERS | _BLOCK_PLACEHOLDERS)