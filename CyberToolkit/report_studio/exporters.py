"""
GhostStrike Report Studio - Exporters
=========================================
Markdown -> HTML (same dark theme as lib/evidence.sh's
gs_evidence_export_report, for visual consistency across every HTML surface
GhostStrike produces).

JSON/SARIF are NOT reimplemented here. gs_finding_export_json/
gs_finding_export_sarif in lib/finding_ontology.sh already generate them
correctly (SARIF 2.1.0, in gs_finding_export_sarif's case) -- but both
operate unconditionally on every finding in GS_FINDINGS_DIR, with no
engagement filter and no dedup-exclusion, which doesn't match what a
per-engagement report needs to show. Rather than re-deriving SARIF
generation a second way (exactly the kind of drift this codebase's own
comments repeatedly warn against), the already-scoped, already-deduped
finding list this package loaded is staged into a throwaway temp
directory and the real bash function is invoked against *that* --
reusing the one real implementation, just pointed at the right input.

Copyright (C) 2026 Fouad Ailabouni. All rights reserved.
"""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path

import markdown as _markdown

from .data_sources import ReportData

_CYBERTOOLKIT_DIR = str(Path(__file__).resolve().parent.parent)
if _CYBERTOOLKIT_DIR not in sys.path:
    sys.path.insert(0, _CYBERTOOLKIT_DIR)
from shell_command_builder import find_bash_invocation as _find_bash_invocation  # noqa: E402

_HTML_SHELL = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; }}
    body {{
      font-family: 'Segoe UI', Arial, sans-serif;
      background: #0d1117; color: #c9d1d9; margin: 0; padding: 32px;
      max-width: 900px; margin-left: auto; margin-right: auto;
    }}
    h1 {{ color: #58a6ff; border-bottom: 2px solid #21262d; padding-bottom: 10px; }}
    h2 {{ color: #79c0ff; margin-top: 30px; border-bottom: 1px solid #21262d; padding-bottom: 6px; }}
    h3 {{ color: #a5d6ff; margin-top: 22px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #21262d; padding: 8px 12px; text-align: left; }}
    th {{ background: #161b22; color: #79c0ff; }}
    code, pre {{ background: #161b22; color: #ffa657; padding: 2px 6px; border-radius: 3px; font-family: Consolas, monospace; }}
    pre {{ padding: 12px; overflow-x: auto; }}
    a {{ color: #58a6ff; }}
    ul, ol {{ line-height: 1.6; }}
    strong {{ color: #e6edf3; }}
    hr {{ border: none; border-top: 1px solid #21262d; margin: 24px 0; }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""


def render_html(markdown_text: str, title: str) -> str:
    body = _markdown.markdown(markdown_text, extensions=["tables", "fenced_code"])
    return _HTML_SHELL.format(title=title, body=body)


def export_html(markdown_text: str, title: str, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(render_html(markdown_text, title))
    return out_path


def export_pdf(data: ReportData, markdown_text: str, title: str, out_path: Path) -> Path:
    """Two-tier: WeasyPrint (real HTML+CSS -> PDF, reusing the exact same
    render_html() the HTML export produces, so the PDF matches it) is the
    primary path -- it's a declared dependency (requirements.txt) meant
    for the actual Linux deployment target, where its native
    Pango/GObject dependencies install cleanly via apt. It cannot be
    exercised on every dev machine (e.g. Windows without those system
    libraries), so a pure-Python fallback (fpdf2, zero native deps) keeps
    PDF export genuinely testable everywhere, at the cost of losing the
    Markdown/CSS richness WeasyPrint provides -- it builds a plain
    structured document straight from ReportData, the same source
    export_docx() uses, rather than trying to re-parse HTML with no CSS
    engine to render it."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        import weasyprint
        html_str = render_html(markdown_text, title)
        weasyprint.HTML(string=html_str).write_pdf(str(out_path))
        return out_path
    except Exception as weasyprint_exc:
        try:
            _export_pdf_fpdf2(data, title, out_path)
            return out_path
        except ImportError as fpdf_exc:
            raise RuntimeError(
                f"PDF export unavailable: WeasyPrint failed ({weasyprint_exc}) "
                f"and the fpdf2 fallback is not installed ({fpdf_exc})"
            ) from weasyprint_exc


def _export_pdf_fpdf2(data: ReportData, title: str, out_path: Path) -> None:
    from fpdf import FPDF

    def latin1(text: str) -> str:
        # Core PDF fonts (Helvetica) are Latin-1 only with no embedded
        # Unicode font available in this fallback tier -- degrade rather
        # than crash on a finding title/description containing e.g. an
        # em dash or non-Latin script.
        return (text or "").encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def mc(text: str, h: float = 5) -> None:
        # multi_cell()'s own default (new_x=XPos.RIGHT) leaves the cursor
        # at the right margin after a full-width (w=0) cell, so the next
        # multi_cell call starts with ~zero width left and raises
        # "Not enough horizontal space to render a single character".
        # Forcing new_x back to the left margin every time is what a
        # normal top-to-bottom flowing document actually needs.
        pdf.multi_cell(0, h, latin1(text), new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "B", 18)
    mc(title, 10)
    pdf.set_font("Helvetica", "", 10)
    eng = data.engagement or {}
    mc(f"Client: {eng.get('client', 'N/A')} | Environment: {eng.get('environment', 'N/A')} | "
       f"Total findings: {len(data.findings)}", 6)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "Risk Overview", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for sev, count in sorted(data.severity_counts.items(),
                              key=lambda kv: {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}
                              .get(kv[0], 5)):
        pdf.cell(0, 6, latin1(f"{sev}: {count}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "Findings", new_x="LMARGIN", new_y="NEXT")
    for f in data.findings_sorted():
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 11)
        mc(f"[{f.get('severity', 'INFO')}] {f.get('title', 'Untitled')}", 7)
        pdf.set_font("Helvetica", "", 9)
        target = f.get("target") or {}
        if target:
            mc(f"Target: {target.get('host', '')}"
               f"{':' + str(target['port']) if target.get('port') else ''}")
        if f.get("cve_ids"):
            mc("CVE: " + ", ".join(f["cve_ids"]))
        mc(f.get("description", ""))
        if f.get("remediation"):
            pdf.set_font("Helvetica", "B", 9)
            mc("Remediation:")
            pdf.set_font("Helvetica", "", 9)
            mc(f["remediation"])

    pdf.output(str(out_path))


_DOCX_SEVERITY_COLOR = {
    "CRITICAL": "C0392B", "HIGH": "E74C3C", "MEDIUM": "E67E22",
    "LOW": "F1C40F", "INFO": "3498DB",
}


def export_docx(data: ReportData, title: str, out_path: Path) -> Path:
    """Real DOCX generation via python-docx -- not an HTML-to-DOCX
    conversion (which loses structure), a purpose-built document: cover
    page, executive summary table, then one section per finding. Item 28
    calls DOCX out specifically ('Consulting companies live in Word'),
    so this produces an actual Word document a consultant can open, not
    a renamed HTML file.

    Built directly from ReportData rather than from a variant's rendered
    Markdown, since python-docx has no Markdown parser and re-parsing
    Markdown back into structure would be more fragile than building the
    document straight from the same source data every variant already
    uses."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()

    doc.add_heading(title, level=0)
    eng = data.engagement or {}
    meta = doc.add_paragraph()
    meta.add_run(f"Client: {eng.get('client', 'N/A')}\n").bold = True
    meta.add_run(f"Environment: {eng.get('environment', 'N/A')}\n")
    meta.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n")
    meta.add_run(f"Total findings: {len(data.findings)}")

    doc.add_heading("Risk Overview", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Severity", "Count"
    for sev, count in sorted(data.severity_counts.items(),
                              key=lambda kv: {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}
                              .get(kv[0], 5)):
        row = table.add_row().cells
        row[0].text, row[1].text = sev, str(count)

    doc.add_heading("Findings", level=1)
    for f in data.findings_sorted():
        heading = doc.add_heading(level=2)
        run = heading.add_run(f"[{f.get('severity', 'INFO')}] {f.get('title', 'Untitled')}")
        color_hex = _DOCX_SEVERITY_COLOR.get(f.get("severity", "INFO"), "7F8C8D")
        run.font.color.rgb = RGBColor.from_string(color_hex)

        p = doc.add_paragraph()
        p.add_run("Finding ID: ").bold = True
        p.add_run(f"{f.get('finding_id', f.get('id', ''))}\n")
        target = f.get("target") or {}
        if target:
            p.add_run("Target: ").bold = True
            p.add_run(f"{target.get('host', '')}"
                       f"{':' + str(target['port']) if target.get('port') else ''}\n")
        if f.get("cve_ids"):
            p.add_run("CVE: ").bold = True
            p.add_run(", ".join(f["cve_ids"]) + "\n")
        if f.get("cvss_score") is not None:
            p.add_run("CVSS: ").bold = True
            p.add_run(f"{f['cvss_score']}\n")

        doc.add_paragraph(f.get("description", ""))

        if f.get("impact"):
            doc.add_heading("Impact", level=3)
            doc.add_paragraph(f["impact"])
        if f.get("remediation"):
            doc.add_heading("Remediation", level=3)
            doc.add_paragraph(f["remediation"])
        if f.get("reproduction_steps"):
            doc.add_heading("Reproduction Steps", level=3)
            for i, step in enumerate(f["reproduction_steps"], 1):
                doc.add_paragraph(f"{i}. {step}")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by GhostStrike Report Studio")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("999999")

    doc.save(str(out_path))
    return out_path


def export_json(data: ReportData, out_path: Path) -> Path:
    """Serializes the already-loaded, already-scoped finding list directly
    -- this list already came from engagement_query.py, the one canonical
    engagement-scoped/dedup-aware loader, so there is nothing to
    re-derive here."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    export = {
        "exported_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "tool": "GhostStrike Report Studio",
        "engagement_id": data.engagement_id,
        "total": len(data.findings),
        "findings": data.findings,
    }
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(export, f, indent=2)
        f.write("\n")
    return out_path


def _scripts_dir() -> Path:
    return Path(__file__).resolve().parent.parent.parent / "bash_scripts_for_pentest"


def export_sarif(data: ReportData, out_path: Path) -> Path:
    """Stages data.findings into a throwaway temp dir and invokes the real
    gs_finding_export_sarif (lib/finding_ontology.sh) against it, so SARIF
    2.1.0 generation stays a single implementation shared with every bash
    caller instead of a second, Python-side one that could drift."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    scripts_dir = _scripts_dir()
    finding_ontology = scripts_dir / "lib" / "finding_ontology.sh"
    if not finding_ontology.exists():
        raise FileNotFoundError(f"lib/finding_ontology.sh not found at {finding_ontology}")

    with tempfile.TemporaryDirectory(prefix="gs_report_sarif_") as tmpdir:
        tmp_findings_dir = Path(tmpdir) / "findings"
        tmp_findings_dir.mkdir()
        for f in data.findings:
            fid = f.get("finding_id") or f"finding_{data.findings.index(f)}"
            with open(tmp_findings_dir / f"{fid}.json", "w", encoding="utf-8") as fh:
                json.dump(f, fh)

        shim = Path(_CYBERTOOLKIT_DIR).parent / "bash_scripts_for_pentest" / "lib" / "_call_lib_function.sh"
        if not shim.exists():
            raise FileNotFoundError(f"lib/_call_lib_function.sh not found at {shim}")

        cmd = _find_bash_invocation(
            str(shim),
            args=[str(finding_ontology), "gs_finding_export_sarif"],
            env_vars={"GS_FINDINGS_DIR": str(tmp_findings_dir)},
            wsl_path_env_keys={"GS_FINDINGS_DIR"},
            path_arg_indices={0},
        )
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        except Exception as exc:
            raise RuntimeError(f"could not invoke gs_finding_export_sarif: {exc}") from exc

        if result.returncode != 0:
            raise RuntimeError(f"gs_finding_export_sarif failed: {result.stderr or result.stdout}")

        sarif_path = tmp_findings_dir / "findings.sarif.json"
        if not sarif_path.exists():
            raise RuntimeError("gs_finding_export_sarif ran but produced no findings.sarif.json")
        shutil.copy(sarif_path, out_path)
        return out_path